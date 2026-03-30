"""
Generate a properly formatted .docx resume from plain text.
Uses python-docx — Word-compatible, ATS-friendly format.

Many company portals and ATS systems (Workday, Greenhouse, Lever)
require .docx over PDF. This generates both.
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_horizontal_line(paragraph):
    """Add a thin horizontal rule under a paragraph (section divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")   # blue line — professional look
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Section detection helpers ─────────────────────────────────────────────────

SECTION_HEADERS = {
    "PROFESSIONAL SUMMARY", "TECHNICAL SKILLS", "PROFESSIONAL EXPERIENCE",
    "EXPERIENCE", "PROJECTS", "EDUCATION", "CERTIFICATIONS", "SKILLS",
    "SUMMARY", "WORK EXPERIENCE",
}

def _is_section_header(line: str) -> bool:
    return line.strip().upper() in SECTION_HEADERS or (
        line.strip().isupper() and len(line.strip()) > 3
    )

def _is_bullet(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("•") or stripped.startswith("-") or stripped.startswith("*")

def _is_contact_line(line: str) -> bool:
    return any(x in line for x in ["@", "linkedin", "github", "|", "http", "://"]) or \
           bool(re.search(r"\d{3}[-.\s]\d{3}[-.\s]\d{4}", line))

def _is_job_title_company(line: str) -> bool:
    """Lines like 'Java Full Stack Developer' or 'SPS Commerce - Minneapolis, MN'."""
    return (
        not _is_bullet(line)
        and not _is_section_header(line)
        and not _is_contact_line(line)
        and len(line.strip()) > 0
    )


# ── Main generator ────────────────────────────────────────────────────────────

def generate_docx(resume_text: str, output_path: str = None,
                  template: str = "classic") -> str:
    """
    Convert plain-text resume to a clean .docx file.
    template: "classic" | "modern" | "executive"
    Returns the output file path.
    """
    if output_path is None:
        os.makedirs("output", exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join("output", f"resume_{timestamp}.docx")

    t = (template or "classic").lower()
    if t == "modern":
        return _generate_modern(resume_text, output_path)
    elif t == "executive":
        return _generate_executive(resume_text, output_path)
    else:
        return _generate_classic(resume_text, output_path)


# ── Classic template (original single-column with blue accents) ───────────────

def _generate_classic(resume_text: str, output_path: str) -> str:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    lines = resume_text.split("\n")
    first_line = True

    for raw_line in lines:
        line = raw_line.rstrip()

        if first_line and line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            _set_font(run, size=18, bold=True, color=(31, 73, 125))
            first_line = False
            continue

        if _is_contact_line(line) and line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            _set_font(run, size=9, color=(89, 89, 89))
            continue

        if _is_section_header(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(line.strip())
            _set_font(run, size=11, bold=True, color=(31, 73, 125))
            _add_horizontal_line(p)
            continue

        if _is_bullet(line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Inches(0.2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            text = line.strip().lstrip("•-* ").strip()
            run = p.add_run(text)
            _set_font(run, size=10)
            continue

        if line.strip():
            has_year = bool(re.search(r"\b20\d{2}\b", line))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4) if has_year else Pt(6)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(line.strip())
            _set_font(run, size=10, bold=(not has_year),
                      color=(64, 64, 64) if has_year else None)
            continue

        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    doc.save(output_path)
    return output_path


# ── Modern template (navy header, teal section accents, compact) ──────────────

def _generate_modern(resume_text: str, output_path: str) -> str:
    """Clean modern look: dark navy name, teal section underlines, tight spacing."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin   = Inches(0.80)
        section.right_margin  = Inches(0.80)

    lines = resume_text.split("\n")
    first_line = True

    for raw_line in lines:
        line = raw_line.rstrip()

        # Name
        if first_line and line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip().upper())
            _set_font(run, name="Calibri Light", size=22, bold=False,
                      color=(20, 40, 80))
            p.paragraph_format.space_after = Pt(2)
            first_line = False
            continue

        # Contact
        if _is_contact_line(line) and line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            _set_font(run, name="Calibri", size=9, color=(80, 80, 80))
            p.paragraph_format.space_after = Pt(1)
            continue

        # Section headers — teal bottom-border, no background fill
        if _is_section_header(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(line.strip())
            _set_font(run, name="Calibri", size=11, bold=True,
                      color=(0, 112, 120))   # teal
            # Bottom border line in teal
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "8")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "007078")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # Bullets
        if _is_bullet(line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Inches(0.15)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            text = line.strip().lstrip("•-* ").strip()
            run = p.add_run(text)
            _set_font(run, name="Calibri", size=10)
            continue

        # Job titles / company / date lines
        if line.strip():
            has_year = bool(re.search(r"\b20\d{2}\b", line))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3) if has_year else Pt(5)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(line.strip())
            _set_font(run, name="Calibri", size=10,
                      bold=(not has_year),
                      color=(100, 100, 100) if has_year else (20, 40, 80))
            continue

        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)

    doc.save(output_path)
    return output_path


# ── Executive template (minimal, spacious, monochrome) ────────────────────────

def _generate_executive(resume_text: str, output_path: str) -> str:
    """Executive look: large centered name, ALL-CAPS bold section headers with thick rule."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    lines = resume_text.split("\n")
    first_line = True

    for raw_line in lines:
        line = raw_line.rstrip()

        # Name — large, centered, black
        if first_line and line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            _set_font(run, name="Georgia", size=24, bold=True,
                      color=(10, 10, 10))
            p.paragraph_format.space_after = Pt(4)
            first_line = False
            continue

        # Contact
        if _is_contact_line(line) and line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            _set_font(run, name="Georgia", size=9, color=(90, 90, 90))
            p.paragraph_format.space_after = Pt(2)
            continue

        # Section headers — ALL CAPS, thick black underline, generous space above
        if _is_section_header(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(line.strip().upper())
            _set_font(run, name="Georgia", size=11, bold=True,
                      color=(10, 10, 10))
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "12")   # thicker than modern
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "0A0A0A")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # Bullets
        if _is_bullet(line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            text = line.strip().lstrip("•-* ").strip()
            run = p.add_run(text)
            _set_font(run, name="Georgia", size=10)
            continue

        # Job title / company / date
        if line.strip():
            has_year = bool(re.search(r"\b20\d{2}\b", line))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5) if has_year else Pt(7)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(line.strip())
            _set_font(run, name="Georgia", size=10,
                      bold=(not has_year),
                      color=(80, 80, 80) if has_year else (10, 10, 10))
            continue

        if not line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

    doc.save(output_path)
    return output_path
